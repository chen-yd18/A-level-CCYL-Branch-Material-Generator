import urllib.request
from bs4 import BeautifulSoup
from docx import Document
from openpyxl import load_workbook
import datetime
import types

# 全局变量：评比材料docx
doc = Document('./test.docx')
table = doc.tables[0]

# 全局变量：班级基本信息
basicXlsx = load_workbook('./基本信息.xlsx')
basicws = basicXlsx['Sheet1']
classID = basicws['B1'].value
gradeID = int(str(classID)[0]) # 不能除10再乘10，要对软院有信心，
                               # 也许以后就壮大到每个年级10个班了
classSize = basicws['B2'].value
branchSize = basicws['B3'].value # 支部人数，一般不等于班级人数
classPioneerCount = basicws['B4'].value
gradePioneerCount = basicws['B5'].value # 全年级积极分子数量
exploreCup = basicws['B6'].value # 求索杯获奖等级
anchorProject = basicws['B7'].value # 固本计划获奖等级
autumnSanitary = basicws['B8'].value
springSanitary = basicws['B9'].value # 平均卫生成绩
volunteerRegisterCount = basicws['B10'].value
volunteerAverageHour = basicws['B11'].value
volunteerParticipantCount = basicws['B12'].value # 志愿
averageGPA = basicws['B13'].value
failCount = basicws['B14'].value
PEfailCount = basicws['B15'].value # 学习成绩相关
campusSportsCount = basicws['B16'].value # 校级体育赛事
deptSportsCount = basicws['B17'].value # 院系级体育赛事
sportsActivityCount = basicws['B18'].value # 群众体育活动
projectTitle = basicws['B19'].value
isResourceSupportProject = basicws['B20'].value
autumnProjectArticleCount = basicws['B21'].value
springProjectArticleCount = basicws['B22'].value # 支部事业
mediaGroupSize = basicws['B23'].value # 宣传小组人数
mediaTitle = basicws['B24'].value
mediaArticleCount = basicws['B25'].value
mediaReadCount = basicws['B26'].value
mediaFocusCount = basicws['B27'].value
mediaTransCount = basicws['B28'].value # 转载次数
mediaShareCount = basicws['B29'].value # 分享转发次数
mediaRewardCount = basicws['B30'].value # 带赞赏文章数
mediaRewardAmount = basicws['B31'].value # 赞赏总金额
studentFestivalShowTitle = basicws['B32'].value # 学生节节目名称
autumnDayParticipantCount = basicws['B33'].value
springDayParticipantCount = basicws['B34'].value # 团日参与人数
autumnDaySummary = basicws['B35'].value
springDaySummary = basicws['B36'].value # 团日总结
basicXlsx.close()

# 全局变量：日期
curYear = datetime.datetime.now().year
curMonth = datetime.datetime.now().month
curDate = datetime.datetime.now().day

# 提取推送的正文
def getText(url):
    response=urllib.request.urlopen(url)
    string=response.read()
    html=string.decode('utf-8')
    bs=BeautifulSoup(html, "html.parser")
    return bs.find(id='js_content').get_text()

# 提取推送的发布日期
def getDate(url):
    response=urllib.request.urlopen(url)
    string=response.read()
    html=string.decode('utf-8')
    pos=html.find('",s="')
    return html[(pos+5):(pos+15)]


# 提取推送作者
def getAuthor(url):
    response=urllib.request.urlopen(url)
    string=response.read()
    html=string.decode('utf-8')
    pos=html.find('<meta property="og:article:author" content="')
    posEnd = pos + 44
    while(html[posEnd] != '"'):
        posEnd = posEnd + 1
    return html[(pos+44):(posEnd)]

# 提取推送标题
def getTitle(url):
    response=urllib.request.urlopen(url)
    string=response.read()
    html=string.decode('utf-8')
    pos=html.find('<meta property="og:title" content="')
    posEnd = pos + 35
    while(html[posEnd] != '"'):
        posEnd = posEnd + 1
    return html[(pos+35):(posEnd)]

# 填充一格内容
def killTodo(r, c, txt):
    for i in range(0,len(table.cell(r,c).paragraphs)):
        if(len(table.cell(r,c).paragraphs[i].runs)>0
           and table.cell(r,c).paragraphs[i].runs[0].text == 'TODO'):
            table.cell(r,c).paragraphs[i].runs[0].text=txt;
    return

# 生成推送完整说明
def articleInfo(link, count):
    out = '《' + getTitle(link) + '》 ' + getDate(link) + ' 截止到 ' + str(curYear) + '-' + str(curMonth) + '-' + str(curDate) + ' 阅读量：' + str(count) + '\n'
    out += '链接：' + link + '\n'
    return out

# 判断T是否为S的子串。如果是，返回T；否则返回None
def check(S,T):
    if(S.find(T)==-1):
        return None
    else:
        return T

# 获取团日环节
def theDaySessions(linkFile,mode):
    # mode==1 秋季
    # mode==2 春季
    # mode==3 秋季+春季
    out = ''
    file = open(linkFile, 'r')
    mapping = dict()
    for eachline in file:
        if(eachline[0]<'0' or eachline[0]>'9'):
            itsYear = int((getDate(eachline))[0:4])
            if (mode==3) or (mode==2 and itsYear==curYear) or (mode==1 and itsYear==curYear-1):
                S = getText(eachline)
                print(S)
                mapping[check(S,'知识竞赛')]=1
                mapping[check(S,'影片欣赏')]=1
                mapping[check(S,'小组讨论')]=1
                mapping[check(S,'分组讨论')]=1
                mapping[check(S,'小组分享')]=1
                mapping[check(S,'分组分享')]=1
                mapping[check(S,'小组交流')]=1
                mapping[check(S,'分组交流')]=1
                mapping[check(S,'影片观赏')]=1
                mapping[check(S,'观影')]=1
                mapping[check(S,'知识竞答')]=1
                mapping[check(S,'嘉宾总结')]=1
                mapping[check(S,'嘉宾分享')]=1
                mapping[check(S,'小组展示')]=1
                mapping[check(S,'分组展示')]=1
                mapping[check(S,'情景剧')]=1
                mapping[check(S,'舞台剧')]=1
                mapping[check(S,'诗朗诵')]=1
                mapping[check(S,'嘉宾演讲')]=1
                mapping[check(S,'主题演讲')]=1
                # blablabla
                
    for key in mapping.keys():
        if key != None:
            if len(out)>0:
                out+='、'
            out += key
    file.close()
    return out

# 获取支部事业活动形式
def projectSessions(mode):
    # mode==False 不包含随团日合办
    # mode==True 包含随团日合办
    out = ''
    xlsx = load_workbook('./支部事业活动.xlsx')
    ws = xlsx['Sheet1']
    mapping = dict()
    row = 2
    while ws.cell(row,1).value != None:
        mapping[ws.cell(row,2).value]=1
        row = row + 1
    for key in mapping.keys():
        if mode or (key != '支部事业随团日合办'):
            if len(out)>0:
                out += '、'
            out += key
    return out
    
# 党的理论学习
def theoryLearning(linkFile):
    ret = ''
    
    # 推送爬取
    file = open(linkFile, 'r')
    mapping = dict()
    for eachline in file:
        author = getAuthor(eachline)
        if mapping.get(author) != None :
            mapping[author].append(getTitle(eachline))
        else:
            mapping[author]=list()
            mapping[author].append(getTitle(eachline))
    for key in mapping.keys():
        ret += '公众号 ' + key + '：\n'
        cnt = 1
        for title in mapping[key]:
            ret += str(cnt) + '、 《' + title + '》\n'
            cnt = cnt + 1
    file.close()

    ret += '统一活动'
    
    # xlsx解析
    xlsx = load_workbook('./党建活动.xlsx')
    ws = xlsx["Sheet1"]
    r = 2
    typeStr = ''
    while ws.cell(row=r,column=1).value != None :
          typeStr = ws.cell(row=r,column=1).value
          outStr = ''
          if typeStr != '其他' :
              outStr += typeStr + ' '
          outStr += ws.cell(row=r,column=2).value + ' '
          outStr += '参与人数：' + str(ws.cell(row=r,column=3).value) + '人'
          ret += outStr + '\n'
          r = r + 1
    xlsx.close()
    return ret

# 主题团日
def theDay(linkFile,mode):
    # mode==0 生成团日材料
    # mode==1 返回秋季学期团日主题
    # mode==2 返回春季学期团日主题
    ret = ''
    
    # 推送爬取
    file = open(linkFile, 'r')
    mapping = dict()
    mappingInfo = dict()
    autumnTitle = 'autumn'
    springTitle = 'spring'
    articleCount = 0
    lastlink = 'last'
    for eachline in file:
        if(eachline[0]>='0' and eachline[0]<='9'):
            mappingInfo[articleCount] = articleInfo(lastlink, int(eachline))
            articleCount = articleCount + 1
        else:
            title = getTitle(eachline)
            lastlink = eachline
            year = int(getDate(eachline)[0:4])
            pos = title.find('|')
            length = len(title)
            if(year==curYear):
                springTitle = title[(pos+1):length]
            else:
                autumnTitle = title[(pos+1):length]
    file.close()
    if mode == 1:
        return autumnTitle
    if mode == 2:
        return springTitle
    
    # 文案生成
    out = '本学年上下学期分别举行了一次团日，按时提交团日策划并根据意见进行修改。\n'
    out += '两次团日主题分别是《' + autumnTitle + '》、《' + springTitle + '》。团日主题结合时事热点，开展' + theDaySessions(linkFile,3) + '等环节。\n'
    out += '参与人数分别为' + str(autumnDayParticipantCount) + '人和' + str(springDayParticipantCount) +'人'
    if autumnDayParticipantCount+springDayParticipantCount >= 0.8*branchSize-0.01:
        out += '，参与度超过4/5。\n'
    else:
        out += '。\n'
    if isResourceSupportProject:
        out += '参选班团资源支持计划并结项。\n'
    out += '微信公众号：' + mediaTitle + '\n'
    out += '团日相关推送：\n'
    for i in range(articleCount):
        out += str(i+1) + '、' + mappingInfo[i]
    out += '团日相关信息以文件形式附'
    return out

# 实践
def practice():
    xlsx = load_workbook('./实践信息.xlsx')
    ws = xlsx['Sheet1']
    row = 2
    out = '\n'
    while ws.cell(row,1).value != None:
        out += ws.cell(row,1).value + ' ' + str(ws.cell(row,2).value) + '人 '
        if ws.cell(row,3).value == '是':
            if ws.cell(row,2).value == 1:
                out += '任支队长 '
            else:
                out += '其中1人任支队长 '
        if ws.cell(row,4).value != None:
            out += '获' + ws.cell(row,4).value
        out += '\n'
        row = row + 1
    xlsx.close()
    return out

# 志愿
def volunteering():
    out = '志愿者注册人数：' + str(volunteerRegisterCount)
    if volunteerRegisterCount >= 0.666*branchSize:
        out += '达到支部总人数2/3以上'
    out += '\n平均志愿时长：' + str(volunteerAverageHour) + 'h\n参与过志愿活动：' + str(volunteerParticipantCount) + '人\n'
    xlsx = load_workbook('./志愿活动.xlsx')
    ws = xlsx['Sheet1']
    row = 2
    while ws.cell(row,1).value != None:
        out += '以软件'
        if ws.cell(row,1).value == '党支部':
            out += str(gradeID)
        else:
            out += str(classID)
        out += ws.cell(row,1).value + '为单位组织' + ws.cell(row,2).value + '志愿活动，参与人数' + str(ws.cell(row,3).value) + '人\n'
        row = row + 1
    if ws['E3'].value != None:
        out += ws['E3'].value
    xlsx.close()
    return out

# 学风建设活动
def learnAtmosphere():
    xlsx = load_workbook('./学风建设活动.xlsx')
    ws = xlsx['Sheet1']
    out = '\n'
    row = 2
    while ws.cell(row,1).value != None:
        out += ws.cell(row,1).value + '参与人数：' + str(ws.cell(row,2).value) + '人\n'
        row = row + 1
    xlsx.close()
    return out

# 科技竞赛
def sciTech():
    xlsx = load_workbook('./科创统计.xlsx')
    ws = xlsx['Sheet1']
    out = '\n'
    row = 2
    while ws.cell(row,1).value != None:
        out += ws.cell(row,1).value + ws.cell(row,2).value + '，' + ws.cell(row,3).value + '；\n'
        row = row + 1
    xlsx.close()
    return out

# SRT
def SRT():
    xlsx = load_workbook('./科创统计.xlsx')
    ws = xlsx['Sheet2']
    out = '\n'
    row = 2
    while ws.cell(row,1).value != None:
        out += ws.cell(row,1).value + '项目“' + ws.cell(row,2).value + '”，' + ws.cell(row,3).value + '；\n'
        row = row + 1
    xlsx.close()
    return out

# 体育赛事
def sports(prizeOnly):
    mapping = dict()
    xlsx = load_workbook('./体育赛事.xlsx')
    ws = xlsx['Sheet1']
    row = 2
    out = '\n'
    while ws.cell(row,1).value != None:
        # 输出体育赛事名称。若该赛事无人取得名次，则在prizeOnly==true时会出BUG！
        if(ws.cell(row,1).value != ws.cell(row-1,1).value):
            out += ws.cell(row,1).value + '：\n'
        
        if (not prizeOnly) or ws.cell(row,4).value != None :
            sstr =  '    ' + ws.cell(row,2).value + '：'
            if isinstance(ws.cell(row,3).value, int) :
                sstr += '共' + str(ws.cell(row,3).value) + '人参加 '
            else:
                sstr += ws.cell(row,3).value
            if ws.cell(row,4).value != None:
                va = ws.cell(row,4).value
                if va == 0:
                    sstr += '获得第1名并创下纪录'
                else:
                    sstr += '获得第' + str(va) + '名'
            sstr += '\n'
            out += sstr
        row = row + 1
    return out

# 体育社团与体育课成绩
def PEGroupGrade():
    xlsx = load_workbook('./群众体育.xlsx')
    ws = xlsx['Sheet1']
    out = '各类体育俱乐部（社团）（'
    count = 0
    row = 2
    while ws.cell(row,1).value != None:
        count = count + 1
        row = row + 1
    out += str(count) + '人次）:\n'
    row = 2
    while ws.cell(row,1).value != None:
        out += '    ' + ws.cell(row,1).value + ' ' + ws.cell(row,2).value + '\n'
        row = row + 1
    if PEfailCount == 0:
        out += '\n无成员体育课不通过'
    else:
        out += '\n' + str(PEfailCount) + '人次体育课不通过'
    return out

# 群众体育活动
def sportsActivity():
    xlsx = load_workbook('./群众体育.xlsx')
    ws = xlsx['Sheet2']
    out = '群众体育活动个人参与次数（共开展次数'
    if sportsActivityCount > 4:
        out += '4次以上）：\n'
    else:
        out += str(sportsActivityCount) + '次）：\n'
    row = 2
    while ws.cell(row,1).value != None:
        out += '    ' + ws.cell(row,1).value + ' ' + ws.cell(row,2).value + str(ws.cell(row,3).value) + '次\n'
        row = row + 1
    return out

# 支部事业参与人次
def projectParticipate():
    xlsx = load_workbook('支部事业活动.xlsx')
    ws = xlsx['Sheet1']
    out = '\n'
    row = 2
    count = 0
    totperson = 0
    while ws.cell(row,1).value != None:
        if ws.cell(row,1).value != ws.cell(row-1,1).value:
            if totperson > 0.8*branchSize+0.001:
                out += '  参与人次超过支部总人数4/5\n'
            out += ws.cell(row,1).value + '学期：\n'
            if ws.cell(row,1).value == '秋季':
                out += '  1、' + str(autumnProjectArticleCount) + '篇' + projectTitle + '相关推送\n'
            else:
                out += '  1、' + str(springProjectArticleCount) + '篇' + projectTitle + '相关推送\n'
            count = 1
            totperson = 0
        count = count + 1
        out += '  ' + str(count) + '、'
        if ws.cell(row,3).value != None:
            out += ws.cell(row,3).value
        out += ws.cell(row,2).value + '，参与人数：' + str(ws.cell(row,4).value) + '人\n'
        totperson = totperson + ws.cell(row,4).value
        row = row + 1
    out += '活动中设有分组展示、分组讨论等环节，增强了集体凝聚力。'
    return out

# 支部事业育人成效
def projectEducate(linkFile):
    out = '1、通过'
    out += projectSessions(False)
    out += ',支部成员在' + projectTitle + '上有了一定的提高\n'
    out += '2、《'
    out += theDay(linkFile,1) + '》主题团日，成员通过'
    out += theDaySessions(linkFile,1)
    out += autumnDaySummary + '。\n'
    out += '3、《'
    out += theDay(linkFile,2) + '》主题团日，成员通过'
    out += theDaySessions(linkFile,2)
    out += springDaySummary + '。\n'
    return out

# 支部事业资源整合
def projectIntegrate(linkFile):
    out = '支部事业以'
    out += projectSessions(True)
    out += '的形式展开。\n'

    # TODO:邀请嘉宾
    mapping = dict()
    out2 = ''
    file = open(linkFile,'r')
    for eachline in file:
        if eachline[0]<'0' or eachline[0]>'9':
            S = getText(eachline)
            mapping[check(S,"班主任")] = 1
            mapping[check(S,"院TMS协会会长")] = 1
            mapping[check(S,"院TMS分会会长")] = 1
            mapping[check(S,"博士生讲师团讲师")] = 1
            mapping[check(S,"博士生讲师团的讲师")] = 1
            mapping[check(S,"博士生讲师")] = 1
            mapping[check(S,"辅导员")] = 1
            mapping[check(S,"校团委干事")] = 1
            mapping[check(S,"院团委书记")] = 1
            mapping[check(S,"院党委书记")] = 1
            mapping[check(S,"校党委书记")] = 1
            mapping[check(S,"校团委书记")] = 1
            mapping[check(S,"特奖获得者")] = 1
            mapping[check(S,"特奖得主")] = 1
    file.close()
    for key in mapping.keys():
        if key!=None:
            if len(out2)>0:
                out2+='、'
            out2+=key
    out2 = '我们邀请了' + out2 + '等来指导我们活动。\n'
    out += out2
    if isResourceSupportProject:
        out += '我们申请了班团资源支持计划，并获得了班团资源支持计划的支持。'
    return out

# 支部事业宣传文章和阅读量
def projectCommunicate(linkFile):
    file = open(linkFile,'r')
    totRead = 0
    for eachline in file:
        if eachline[0]>='0' and eachline[0]<='9':
            totRead += int(eachline)
    file.close()
    out = '总计阅读量：' + str(totRead) + '\n'
    count = 0
    file = open(linkFile,'r')
    lastlink = ' '
    for eachline in file:
        if eachline[0]>='0' and eachline[0]<='9':
            count += 1
            out += str(count) + '、推送：'
            out += articleInfo(lastlink,int(eachline))
        else:
            lastlink = eachline
    return out

# 特色工作
def specialWorks():
    grade = ((curYear-gradeID)%10)+1
    out = ' '
    if grade == 1:
        out = '本学年为支部成员进入大学阶段的第一学年，支部刚刚组建，'
    if grade == 2:
        out = '本学年为支部成员进入大学阶段的第二学年，是支部集体建设非常关键的一年，'
    if grade == 3:
        out = '本学年为支部成员进入大学阶段的第三学年，同学们日渐忙碌，但班团建设不能放松，'
    out += '我们举行了丰富的集体建设活动，增进支部成员间的关系，除了已提到的活动，还有：\n'
    xlsx = load_workbook('./其他活动.xlsx')
    ws = xlsx['Sheet1']
    row = 2
    while ws.cell(row,1).value != None:
        out += '  ' + str(row-1) + '、' + ws.cell(row,1).value + '\n'
        row = row + 1
    if grade == 1:
        out += '在文艺活动方面，支部成员参与了新生舞会、五系歌赛与院学生节。'
    else:
        out += '在文艺活动方面，支部成员参与了五系歌赛与院学生节。'
    if studentFestivalShowTitle != None:
        out += '支部成员总负责并参与了节目' + studentFestivalShowTitle + '。\n'
    out += '在宣传上，我们拥有' + str(mediaGroupSize) + '人的宣传小组团队，主要负责管理我们的公众号“' + mediaTitle
    out += '”，截至' + str(curMonth) + '月' + str(curDate) + '日“' + mediaTitle + '”关注人数' + str(mediaFocusCount)
    out += '人，推送数' + str(mediaArticleCount) + '条，总计阅读量达' + str(mediaReadCount) + '次，平均每篇阅读量'
    out += str(mediaReadCount//mediaArticleCount) + '次。文章被各类账号转载' + str(mediaTransCount) + '次，被分享转发' + str(mediaShareCount)
    out += '次。发表带赞赏文章' + str(mediaRewardCount) + '篇，合计赞赏金额' + str(mediaRewardAmount) + '元。\n'
    out += '“' + mediaTitle + '”' + str(mediaArticleCount) + '条推送包括生日推送、活动推送、学习类推送、分享类推送、学生节节目推送、图片分享……'
    return out

# 主函数
def run():
    killTodo(2,5,theoryLearning('./theoryLearningLinks.txt'))
    print('ok 1/21')
    killTodo(3,5,'求索杯软件学院初赛'+str(gradeID)+'字班取得'+str(exploreCup)+'等奖\n'+str(curYear-1)+'秋软件学院“固本计划”读书分享活动中软件'+str(classID)+'党课小组取得'+str(anchorProject)+'等奖')
    print('ok 2/21')
    killTodo(4,5,theDay('./theDay.txt',0))
    print('ok 3/21')
    killTodo(6,5,practice())
    print('ok 4/21')
    killTodo(7,5,'秋季学期平均成绩 ' + str(autumnSanitary) + '\n春季学期平均成绩 ' + str(springSanitary) + '\n学年平均成绩 ' + str(0.5*(autumnSanitary+springSanitary)))
    print('ok 5/21')
    killTodo(8,5,volunteering())
    print('ok 6/21')
    killTodo(9,5,'均绩' + str(averageGPA) + ' 人均不及格门次：' + str(failCount*1.0/classSize))
    print('ok 7/21')
    killTodo(10,5,learnAtmosphere())
    print('ok 8/21')
    killTodo(11,5,sciTech())
    print('ok 9/21')
    killTodo(12,5,SRT())
    print('ok 10/21')
    killTodo(13,5,sports(True))
    print('ok 11/21')
    killTodo(14,5,'支部成员共参加学校各类体育赛事' + str(campusSportsCount) + '人次，系内赛' + str(deptSportsCount) + '人次\n' + sports(False))
    print('ok 12/21')
    killTodo(15,5,PEGroupGrade())
    print('ok 13/21')
    killTodo(16,5,sportsActivity())
    print('ok 14/21')
    killTodo(18,5,projectParticipate())
    print('ok 15/21')
    Edu = projectEducate('./theDay.txt')
    killTodo(19,5,Edu)
    print('ok 16/21')
    killTodo(20,5,projectIntegrate('./theDay.txt'))
    print('ok 17/21')
    killTodo(21,5,projectCommunicate('./theProject.txt'))
    print('ok 18/21')
    killTodo(22,5,'活动形式丰富，活动有趣且效果良好，支部成员参与度高、收获大，支部事业开展情况及时在公众号上更新宣传。\n' + Edu)
    print('ok 19/21')
    killTodo(23,5,'开展以' + projectTitle + '为主题的支部事业一年。\n' + Edu)
    print('ok 20/21')
    killTodo(24,5,specialWorks())
    print('ok 21/21')
    doc.save('./test_new.docx')
    return

run()
